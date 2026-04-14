#!/usr/bin/env python3
"""
Генерация DOCX: Концепция мобильного приложения для родительского комитета
с интеграцией коллективного счёта Ozon Bank.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

doc = Document()

# ── Стили ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Заголовки
for level, (size, color_hex) in enumerate([
    (Pt(26), '1B3A5C'),  # Heading 1
    (Pt(18), '2E5E8E'),  # Heading 2
    (Pt(14), '3A7BBF'),  # Heading 3
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.size = size
    h.font.color.rgb = RGBColor.from_string(color_hex)
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level > 1 else 24)
    h.paragraph_format.space_after = Pt(8)

# ── Разделы страницы ──────────────────────────────────────
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

# ── Утилиты ───────────────────────────────────────────────
def add_colored_cell(cell, text, bg='D5E8F0', bold=False, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Arial'
    run.bold = bold
    return run

def add_table_row(table, cells_data, header=False, row_idx=0):
    row = table.add_row()
    for i, (text, width) in enumerate(cells_data):
        cell = row.cells[i]
        bg = '1B3A5C' if header else ('F5F8FB' if row_idx % 2 == 0 else 'FFFFFF')
        color = 'FFFFFF' if header else '333333'
        r = add_colored_cell(cell, text, bg=bg, bold=header, size=Pt(9) if not header else Pt(10))
        r.font.color.rgb = RGBColor.from_string(color)

def add_bullet(doc, text, level=0, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = 'Arial'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    return p

def add_numbered(doc, text, bold_prefix=''):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = 'Arial'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string('666666')
    run.font.name = 'Arial'

# ══════════════════════════════════════════════════════════
# ТИТУЛЬНАЯ СТРАНИЦА
# ══════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('РОДИТЕЛЬСКИЙ КОМИТЕТ')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor.from_string('1B3A5C')
run.font.name = 'Arial'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Мобильное приложение для прозрачного\nуправления финансами класса')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string('4A7FB5')
run.font.name = 'Arial'

doc.add_paragraph()

# Мета-информация
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime('%d.%m.%Y')
run = meta.add_run(f'Концепция и план реализации\nВерсия 1.0 | {today}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string('888888')
run.font.name = 'Arial'

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# ОГЛАВЛЕНИЕ
# ══════════════════════════════════════════════════════════
doc.add_heading('Оглавление', level=1)
toc_items = [
    ('1. Проблема и контекст', 3),
    ('2. Решение: обзор продукта', 3),
    ('3. Интеграция с Ozon Bank (коллективный счёт)', 4),
    ('4. Функциональные требования (MVP)', 5),
    ('5. Архитектура и технологический стек', 7),
    ('6. Дизайн и UX', 8),
    ('7. Юридические аспекты (РФ)', 9),
    ('8. Дорожная карта (Roadmap)', 10),
    ('9. Монетизация и масштабирование', 11),
    ('10. Риски и митигация', 12),
    ('11. Критика и рекомендации автора', 13),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{item} ')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    # dots
    run2 = p.add_run('.' * (70 - len(item)))
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor.from_string('CCCCCC')
    run3 = p.add_run(f' {page}')
    run3.font.size = Pt(11)
    run3.font.name = 'Arial'

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. ПРОБЛЕМА И КОНТЕКСТ
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Проблема и контекст', level=1)

doc.add_heading('1.1. Текущая ситуация', level=2)
p = doc.add_paragraph(
    'Родительские комитеты в российских школах ежегодно собирают средства на нужды класса: '
    'подарки учителям, экскурсии, ремонт, выпускные мероприятия и прочее. '
    'Типичный процесс выглядит так:'
)
p.style.font.name = 'Arial'

add_bullet(doc, 'Казначей (обычно один родитель) собирает наличные или переводы на личную карту')
add_bullet(doc, 'Ведёт учёт в тетради, Excel или заметках в телефоне')
add_bullet(doc, 'Периодически отчитывается в родительском чате (WhatsApp/Telegram)')
add_bullet(doc, 'Родители не видят реальных транзакций, доверие основано на честном слове')

doc.add_heading('1.2. Ключевые боли', level=2)

# Таблица болей
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Ширины колонок
widths = [Cm(3.5), Cm(6), Cm(6.5)]
for i, width in enumerate(widths):
    table.columns[i].width = width

# Заголовок
for i, text in enumerate(['Проблема', 'Описание', 'Последствия']):
    cell = table.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

pain_data = [
    ('Блокировка карт', 'Банки блокируют карты при множественных однотипных переводах от разных лиц (115-ФЗ)', 'Невозможность расходования средств, стресс'),
    ('Отсутствие прозрачности', 'Родители не видят реальное движение денег', 'Недоверие, конфликты в чате'),
    ('Ручной учёт', 'Excel/тетрадь — ошибки, потеря данных', 'Несоответствие баланса, обвинения'),
    ('Налоговые риски', 'Переводы на личную карту могут быть квалифицированы как доход', 'Штрафы, доначисления НДФЛ'),
    ('Время казначея', 'Сверка, отчёты, напоминания — часы работы', 'Выгорание, отказ от роли'),
]
for ri, (prob, desc, cons) in enumerate(pain_data):
    row = table.add_row()
    for i, text in enumerate([prob, desc, cons]):
        cell = row.cells[i]
        cell.width = widths[i]
        bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
        add_colored_cell(cell, text, bg=bg, size=Pt(9))

doc.add_paragraph()

doc.add_heading('1.3. Почему коллективный счёт Ozon Bank — это game changer', level=2)
p = doc.add_paragraph(
    'Ozon Bank запустил продукт «Коллективный счёт», который решает '
    'фундаментальную проблему — средства больше не поступают на личную карту одного человека. '
    'Ключевые возможности:'
)

add_bullet(doc, 'Единый счёт ', bold_prefix='Юридически отделён от личных средств. ')
add_bullet(doc, 'Настраиваемые права — кто может пополнять, кто тратить, кто только просматривать')
add_bullet(doc, 'Банковские карты к счёту — возможность оплаты напрямую со счёта')
add_bullet(doc, 'История транзакций — все операции видны участникам')
add_bullet(doc, 'Без блокировок — счёт создан именно для коллективных сборов')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. РЕШЕНИЕ
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Решение: обзор продукта', level=1)

doc.add_heading('2.1. Миссия', level=2)
p = doc.add_paragraph()
run = p.add_run('Сделать финансы родительского комитета полностью прозрачными, '
                'автоматизированными и юридически чистыми.')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Arial'

doc.add_heading('2.2. Целевая аудитория', level=2)
add_bullet(doc, 'Казначеи родительских комитетов (основной пользователь)', bold_prefix='Первичная: ')
add_bullet(doc, 'Родители учеников (просмотр, пополнение)', bold_prefix='Вторичная: ')
add_bullet(doc, 'Классные руководители (опционально, только просмотр)', bold_prefix='Третичная: ')

doc.add_heading('2.3. Ключевое ценностное предложение', level=2)

# Таблица ценностей
table2 = doc.add_table(rows=1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'

for i, text in enumerate(['Для казначея', 'Для родителей', 'Для класса']):
    cell = table2.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

row = table2.add_row()
texts = [
    'Автоматический учёт.\nОтчёты в 1 клик.\nНет личной ответственности за средства.',
    'Полная прозрачность.\nЖивые операции в реальном времени.\nДоверие через факты, не слова.',
    'Юридическая чистота.\nНикаких блокировок.\nИстория сохраняется годами.'
]
for i, text in enumerate(texts):
    add_colored_cell(row.cells[i], text, bg='EAF2F8', size=Pt(9))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. ИНТЕГРАЦИЯ С OZON BANK
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Интеграция с Ozon Bank (коллективный счёт)', level=1)

doc.add_heading('3.1. Модель интеграции', level=2)
p = doc.add_paragraph(
    'На начальном этапе (MVP) приложение не будет напрямую интегрироваться с API Ozon Bank. '
    'Вместо этого используется гибридная модель:'
)

add_numbered(doc, 'Казначей создаёт коллективный счёт в приложении Ozon Bank вручную', bold_prefix='Шаг 1: ')
add_numbered(doc, 'В нашем приложении создаётся «Класс» и привязывается номер счёта', bold_prefix='Шаг 2: ')
add_numbered(doc, 'Родители добавляются в класс (по ссылке-приглашению или QR-коду)', bold_prefix='Шаг 3: ')
add_numbered(doc, 'Родители пополняют счёт через Ozon Bank / СБП по реквизитам из приложения', bold_prefix='Шаг 4: ')
add_numbered(doc, 'Казначей вносит операции (расходы) в приложение вручную или через фото чека', bold_prefix='Шаг 5: ')

doc.add_heading('3.2. Эволюция интеграции', level=2)

table3 = doc.add_table(rows=1, cols=3)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.style = 'Table Grid'

for i, text in enumerate(['Этап', 'Интеграция', 'Описание']):
    cell = table3.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

stages = [
    ('MVP', 'Ручная', 'Казначей вносит данные. Фото чеков с OCR-распознаванием'),
    ('v2.0', 'Полуавтоматическая', 'Импорт выписки банка (CSV/PDF). Автоматическое сопоставление'),
    ('v3.0', 'Open API', 'Прямое подключение к API банка (при наличии). Автоматическая синхронизация'),
]
for ri, (stage, integ, desc) in enumerate(stages):
    row = table3.add_row()
    for i, text in enumerate([stage, integ, desc]):
        bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
        add_colored_cell(row.cells[i], text, bg=bg, size=Pt(9))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. ФУНКЦИОНАЛЬНЫЕ ТРЕБОВАНИЯ (MVP)
# ══════════════════════════════════════════════════════════
doc.add_heading('4. Функциональные требования (MVP)', level=1)

doc.add_heading('4.1. Роли пользователей', level=2)
roles = [
    ('Администратор (казначей)', [
        'Создание/управление классом',
        'Добавление/удаление участников',
        'Создание сборов (целей)',
        'Внесение расходов с прикреплением чеков',
        'Формирование отчётов',
        'Рассылка уведомлений',
    ]),
    ('Родитель', [
        'Просмотр баланса и истории операций',
        'Просмотр активных сборов',
        'Подтверждение своего взноса',
        'Просмотр отчётов и чеков',
        'Получение уведомлений',
    ]),
    ('Наблюдатель (классный руководитель)', [
        'Только просмотр баланса и отчётов',
        'Без доступа к персональным данным родителей',
    ]),
]

for role, perms in roles:
    doc.add_heading(role, level=3)
    for perm in perms:
        add_bullet(doc, perm)

doc.add_heading('4.2. Основные экраны MVP', level=2)

screens = [
    ('Главный экран', 'Текущий баланс счёта, активные сборы, последние операции, быстрые действия'),
    ('Сбор средств', 'Создание цели (название, сумма, срок). Прогресс-бар. Список кто сдал / не сдал. Мягкие напоминания'),
    ('Операции', 'Лента транзакций (пополнения + расходы). Фильтры по типу/дате. Прикреплённые чеки/фото'),
    ('Отчёт', 'Автогенерация отчёта за период. Диаграмма расходов по категориям. Экспорт в PDF/отправка в чат'),
    ('Класс', 'Список участников и их роли. Приглашение по ссылке/QR. Настройки прав'),
    ('Профиль', 'Имя, контактные данные. Настройки уведомлений. Привязанные классы'),
]

table4 = doc.add_table(rows=1, cols=2)
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.style = 'Table Grid'

for i, text in enumerate(['Экран', 'Функционал']):
    cell = table4.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for ri, (screen, func) in enumerate(screens):
    row = table4.add_row()
    bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
    r = add_colored_cell(row.cells[0], screen, bg=bg, bold=True, size=Pt(9))
    add_colored_cell(row.cells[1], func, bg=bg, size=Pt(9))

doc.add_heading('4.3. Ключевые фичи', level=2)
add_bullet(doc, 'Автоматическое распознавание суммы и магазина из фото чека (OCR)', bold_prefix='OCR чеков: ')
add_bullet(doc, 'Max Messenger бот (основной) + Telegram-бот (параллельный) + Push для нативного приложения (v2)', bold_prefix='Уведомления: ')
add_bullet(doc, 'Генерация QR-кода с реквизитами для пополнения коллективного счёта', bold_prefix='QR-оплата: ')
add_bullet(doc, 'Автоматическая пометка «Не сдал» с мягким напоминанием', bold_prefix='Статус взносов: ')
add_bullet(doc, 'Круговая диаграмма расходов, история по месяцам', bold_prefix='Аналитика: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. АРХИТЕКТУРА
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Архитектура и технологический стек', level=1)

doc.add_heading('5.1. Рекомендуемый стек', level=2)

table5 = doc.add_table(rows=1, cols=3)
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
table5.style = 'Table Grid'

for i, text in enumerate(['Компонент', 'Технология', 'Обоснование']):
    cell = table5.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

stack = [
    ('Боты (MVP)', 'Max Messenger Bot + Telegram Bot', 'Нулевой порог входа — родителям не нужно ставить приложение'),
    ('Mobile (v2)', 'Flutter', 'Кроссплатформа: iOS + Android из одной кодовой базы'),
    ('Backend', 'Node.js (NestJS) или Python (FastAPI)', 'Быстрая разработка, хорошее комьюнити в РФ'),
    ('База данных', 'PostgreSQL + Redis', 'Надёжность, JSON-поддержка, кэширование'),
    ('Хранение файлов', 'S3-совместимое (Yandex Object Storage)', 'Чеки и фото в облаке РФ'),
    ('OCR', 'Yandex Vision или Google ML Kit', 'Распознавание чеков, on-device для скорости'),
    ('Push', 'Firebase Cloud Messaging', 'Бесплатно, надёжно, кроссплатформа'),
    ('Авторизация', 'SMS OTP (через SMS.ru / SMSC)', 'Привычно для аудитории, без паролей'),
    ('Хостинг', 'Yandex Cloud / Selectel', 'Данные в РФ (152-ФЗ), приемлемые цены'),
]
for ri, (comp, tech, reason) in enumerate(stack):
    row = table5.add_row()
    bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
    add_colored_cell(row.cells[0], comp, bg=bg, bold=True, size=Pt(9))
    add_colored_cell(row.cells[1], tech, bg=bg, size=Pt(9))
    add_colored_cell(row.cells[2], reason, bg=bg, size=Pt(9))

doc.add_heading('5.2. Архитектурная схема (упрощённая)', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '  [Max Bot / TG Bot / Flutter (v2)]  <-->  [API Gateway / NestJS]\n'
    '                            |\n'
    '              +-------------+-------------+\n'
    '              |             |             |\n'
    '        [PostgreSQL]  [Redis Cache]  [S3 Storage]\n'
    '              |                           |\n'
    '        [OCR Service]              [Чеки/Фото]\n'
    '              |\n'
    '     [Push / Max Bot / TG Bot]'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor.from_string('333333')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. ДИЗАЙН И UX
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Дизайн и UX', level=1)

doc.add_heading('6.1. Принципы', level=2)
add_numbered(doc, 'Мама с ребёнком на руках должна разобраться за 30 секунд', bold_prefix='Простота: ')
add_numbered(doc, 'Все суммы, чеки и операции видны каждому участнику', bold_prefix='Прозрачность: ')
add_numbered(doc, 'Мягкие формулировки, без агрессии. «Напомнить» вместо «Потребовать»', bold_prefix='Деликатность: ')
add_numbered(doc, 'Оффлайн-просмотр, загрузка чеков при появлении сети', bold_prefix='Оффлайн-first: ')

doc.add_heading('6.2. Цветовая палитра', level=2)
colors_table = doc.add_table(rows=1, cols=3)
colors_table.style = 'Table Grid'
for i, text in enumerate(['Цвет', 'Hex', 'Назначение']):
    cell = colors_table.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

palette = [
    ('Основной', '#2E75B6', 'Кнопки, акценты, заголовки'),
    ('Успех', '#27AE60', 'Пополнения, подтверждения'),
    ('Предупреждение', '#F39C12', 'Напоминания, ожидание'),
    ('Фон', '#F8FAFB', 'Основной фон приложения'),
    ('Текст', '#2C3E50', 'Основной текст'),
]
for ri, (name, hex_c, purpose) in enumerate(palette):
    row = colors_table.add_row()
    bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
    add_colored_cell(row.cells[0], name, bg=bg, bold=True, size=Pt(9))
    add_colored_cell(row.cells[1], hex_c, bg=bg, size=Pt(9))
    add_colored_cell(row.cells[2], purpose, bg=bg, size=Pt(9))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. ЮРИДИЧЕСКИЕ АСПЕКТЫ
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Юридические аспекты (РФ)', level=1)

doc.add_heading('7.1. Регуляторные требования', level=2)
add_bullet(doc, 'Хранение персональных данных граждан РФ на территории РФ. Использование российских облаков', bold_prefix='152-ФЗ (Персональные данные): ')
add_bullet(doc, 'Коллективный счёт банка снимает риск квалификации переводов как доход физлица', bold_prefix='115-ФЗ (ПОД/ФТ): ')
add_bullet(doc, 'Приложение НЕ является платёжным агентом — деньги идут напрямую на счёт банка', bold_prefix='Платёжный агент: ')
add_bullet(doc, 'Сбор школой/комитетом — добровольные пожертвования. Необходима оферта', bold_prefix='Налогообложение: ')

doc.add_heading('7.2. Необходимые документы', level=2)
add_numbered(doc, 'Политика конфиденциальности (обязательна для App Store / Google Play)')
add_numbered(doc, 'Пользовательское соглашение (оферта)')
add_numbered(doc, 'Согласие на обработку персональных данных')
add_numbered(doc, 'Положение о родительском комитете класса (шаблон для школ)')

add_note(doc, 'Рекомендация: проконсультироваться с юристом до запуска даже в бета-режиме.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. ДОРОЖНАЯ КАРТА
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Дорожная карта (Roadmap)', level=1)

roadmap = doc.add_table(rows=1, cols=4)
roadmap.alignment = WD_TABLE_ALIGNMENT.CENTER
roadmap.style = 'Table Grid'

for i, text in enumerate(['Этап', 'Срок', 'Функционал', 'Цель']):
    cell = roadmap.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

phases = [
    ('0. Прототип', '2-3 недели', 'Figma-макеты, UX-тесты с 3-5 родителями из класса', 'Валидация идеи'),
    ('1. MVP', '6-8 недель', 'Регистрация, создание класса, ручной ввод операций, просмотр баланса, базовые отчёты', 'Обкатка на своём классе (25-30 родителей)'),
    ('2. Бета', '4-6 недель', 'OCR чеков, push-уведомления, экспорт PDF, Telegram-бот параллельно с Max', 'Расширение на 3-5 классов в школе'),
    ('3. Публичный запуск', '4-6 недель', 'App Store + Google Play, мультишкольность, импорт выписок', 'Первые 50 классов'),
    ('4. Масштабирование', '8-12 недель', 'API-интеграция с банками, аналитика, монетизация', 'Рост по школам РФ'),
]
for ri, (phase, term, func, goal) in enumerate(phases):
    row = roadmap.add_row()
    bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
    add_colored_cell(row.cells[0], phase, bg=bg, bold=True, size=Pt(9))
    add_colored_cell(row.cells[1], term, bg=bg, size=Pt(9))
    add_colored_cell(row.cells[2], func, bg=bg, size=Pt(9))
    add_colored_cell(row.cells[3], goal, bg=bg, size=Pt(9))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 9. МОНЕТИЗАЦИЯ
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Монетизация и масштабирование', level=1)

doc.add_heading('9.1. Модели монетизации', level=2)
add_bullet(doc, 'Бесплатно для 1 класса (до 35 родителей). Базовый функционал', bold_prefix='Freemium: ')
add_bullet(doc, '299-499 руб/мес за класс — расширенная аналитика, автоотчёты, OCR без лимита, приоритетная поддержка', bold_prefix='Premium: ')
add_bullet(doc, '990-2990 руб/мес — мультиклассы, брендирование школы, админ-панель для завуча', bold_prefix='Школа: ')
add_bullet(doc, 'Партнёрство с банками за привлечение клиентов на коллективные счета (CPA)', bold_prefix='B2B (банки): ')

doc.add_heading('9.2. Потенциал рынка', level=2)
p = doc.add_paragraph()
run = p.add_run('Россия:')
run.bold = True
run.font.name = 'Arial'
add_bullet(doc, '~42 000 школ, ~16.5 млн учеников, ~500 000 классов')
add_bullet(doc, 'При конверсии 5% premium = 25 000 классов x 399 руб/мес = ~10 млн руб/мес')

doc.add_heading('9.3. Международный потенциал', level=2)
add_bullet(doc, 'Казахстан, Беларусь, Узбекистан — аналогичная модель родительских комитетов')
add_bullet(doc, 'Европа — PTA (Parent-Teacher Association) с аналогичными проблемами')
add_bullet(doc, 'Адаптация: локализация, интеграция с местными банками/платёжными системами')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 10. РИСКИ
# ══════════════════════════════════════════════════════════
doc.add_heading('10. Риски и митигация', level=1)

risks_table = doc.add_table(rows=1, cols=4)
risks_table.alignment = WD_TABLE_ALIGNMENT.CENTER
risks_table.style = 'Table Grid'

for i, text in enumerate(['Риск', 'Вероятность', 'Влияние', 'Митигация']):
    cell = risks_table.rows[0].cells[i]
    add_colored_cell(cell, text, bg='1B3A5C', bold=True, size=Pt(10))
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

risks = [
    ('Ozon Bank закроет продукт', 'Низкая', 'Высокое', 'Архитектура bank-agnostic. Поддержка любого банка с коллективными счетами'),
    ('Низкое adoption', 'Средняя', 'Высокое', 'Обкатка на своём классе. Product-led growth. Max + Telegram боты — максимальный охват'),
    ('Регуляторные изменения', 'Средняя', 'Среднее', 'Юрист на retainer. Гибкая архитектура'),
    ('Конкуренция', 'Средняя', 'Среднее', 'Фокус на UX и «бесшовность». Быстрое движение'),
    ('Утечка данных', 'Низкая', 'Высокое', 'Шифрование, минимизация данных, аудит безопасности'),
]
for ri, (risk, prob, impact, mitigation) in enumerate(risks):
    row = risks_table.add_row()
    bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
    add_colored_cell(row.cells[0], risk, bg=bg, bold=True, size=Pt(9))
    add_colored_cell(row.cells[1], prob, bg=bg, size=Pt(9))
    add_colored_cell(row.cells[2], impact, bg=bg, size=Pt(9))
    add_colored_cell(row.cells[3], mitigation, bg=bg, size=Pt(9))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 11. КРИТИКА И РЕКОМЕНДАЦИИ
# ══════════════════════════════════════════════════════════
doc.add_heading('11. Критика и рекомендации автора', level=1)

doc.add_heading('11.1. Что я бы добавил от себя', level=2)

add_bullet(doc, 'Не делайте полноценное мобильное приложение сразу. '
    'Начните с бота в Max Messenger — это бесплатная дистрибуция, '
    'нулевая стоимость установки. Параллельно запустите Telegram-бота для охвата. '
    'Если product-market fit подтвердится — тогда нативное приложение на Flutter.',
    bold_prefix='Max Messenger бот вместо нативного MVP: ')

add_bullet(doc, 'Добавьте возможность для каждого родителя видеть '
    'именно свою историю взносов: сколько сдал, за что, когда. '
    'Это снимает 90% вопросов в чате.',
    bold_prefix='«Личный кабинет родителя»: ')

add_bullet(doc, 'После каждого расхода автоматически отправлять в общий чат '
    'карточку: «Потрачено 3 500 руб — Букеты учителям (фото чека)». '
    'Это создаёт эффект «живой отчётности».',
    bold_prefix='Автоматический отчёт в чат: ')

add_bullet(doc, 'Подумайте о сценарии, когда родитель не может заплатить '
    'полную сумму сразу. Разбивка на 2-3 платежа снижает барьер.',
    bold_prefix='Рассрочка взносов: ')

add_bullet(doc, 'Добавьте опциональное голосование: «Покупаем подарок за 5000 или за 3000?» '
    'Это снимает конфликты до того, как они возникнут.',
    bold_prefix='Голосование по тратам: ')

add_bullet(doc, 'Казначеи выгорают за 1-2 года. Сделайте передачу дел '
    'максимально простой: нажал кнопку — вся история, права и настройки перешли к новому человеку.',
    bold_prefix='«Передача дел»: ')

doc.add_heading('11.2. О чём стоит подумать заранее', level=2)

add_bullet(doc, 'Будут ситуации, когда один родитель не сдаёт деньги месяцами. '
    'Приложение не должно превращаться в инструмент давления. '
    'Деликатные напоминания — да. Публичный «позор» — нет.',
    bold_prefix='Конфликтные сценарии: ')

add_bullet(doc, 'Если в классе 30 родителей — это 30 потенциальных субъектов 152-ФЗ. '
    'Минимизируйте данные: имя + телефон, никаких паспортов и адресов.',
    bold_prefix='Персональные данные: ')

add_bullet(doc, 'Ozon Bank может быть не единственным вариантом. '
    'Тинькофф, Сбербанк, Альфа — все могут запустить аналоги. '
    'Архитектура должна быть bank-agnostic с первого дня.',
    bold_prefix='Зависимость от одного банка: ')

doc.add_heading('11.3. Первый шаг прямо сейчас', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'Создайте коллективный счёт в Ozon Bank, добавьте 5 самых активных родителей, '
    'и проведите один сбор полностью через него. Без приложения. Вручную. '
    'Это даст вам понимание реальных проблем и ответ на вопрос — '
    'нужно ли приложение или достаточно ботов в Max и Telegram.'
)
run.font.size = Pt(11)
run.font.name = 'Arial'
run.italic = True

doc.add_paragraph()

# Финальная подпись
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'Документ подготовлен: {today}')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor.from_string('999999')
run.font.name = 'Arial'

# ── Колонтитулы ──────────────────────────────────────────
from docx.oxml.ns import qn
for section in doc.sections:
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run('Родительский Комитет | Концепция приложения')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string('AAAAAA')
    run.font.name = 'Arial'

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('Конфиденциально')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string('CCCCCC')
    run.font.name = 'Arial'

# ── Сохранение ────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output', 'Концепция_Родительский_Комитет.docx')
doc.save(output_path)
print(f'OK: {output_path}')
